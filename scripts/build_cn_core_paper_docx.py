"""Build the FADE-Net Chinese journal manuscript as a styled DOCX.

Design system: ``narrative_proposal`` with an academic A4 override.
Header pattern: ``memo_masthead`` on the first page, compact running header
afterwards.  The source of truth for prose and data remains the Markdown file.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "paper" / "FADE-Net_中文核心论文初稿.md"
OUTPUT = ROOT / "output" / "doc" / "FADE-Net_中文核心论文优化稿.docx"
DOCUMENT_SKILL_SCRIPTS = Path(
    "C:/Users/Administrator/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.805.11740/skills/documents/scripts"
)
if str(DOCUMENT_SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(DOCUMENT_SKILL_SCRIPTS))

from table_geometry import (  # noqa: E402
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)


PRESET = "narrative_proposal"
HEADER_PATTERN = "memo_masthead"
TOKENS = {
    "page": {"size": "A4", "margin_top_cm": 2.2, "margin_bottom_cm": 2.1, "margin_side_cm": 2.4},
    "font": {"zh_body": "宋体", "zh_heading": "黑体", "latin": "Times New Roman", "math": "Cambria Math"},
    "size_pt": {"body": 10.5, "title": 18, "subtitle": 11.5, "h1": 13, "h2": 11, "caption": 9, "table": 8.2, "reference": 9},
    "color": {"ink": "20262E", "accent": "4E708C", "muted": "65717D", "rule": "AEB8C2", "header_fill": "F3F5F7"},
    "spacing_pt": {"body_after": 3, "h1_before": 12, "h1_after": 6, "h2_before": 8, "h2_after": 4},
}

CITATIONS_AS_SUPERSCRIPT = False
CITATION_TOKEN = r"\[(?:\d+(?:-\d+)?)(?:(?:,|，)\s*\d+(?:-\d+)?)*\]"
INLINE_TOKEN_RE = re.compile(rf"(`[^`]+`|\$[^$]+\$|{CITATION_TOKEN})")
TABLE_SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")
FIGURE_RE = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<path>[^)]+)\)$")
NUMBERED_RE = re.compile(r"^(?P<number>\d+)\.\s+(?P<text>.+)$")
REFERENCE_RE = re.compile(r"^\[(?P<number>\d+)\]\s+(?P<text>.+)$")

EQUATION_LINEAR = {
    1: "qₖ(y) = exp[−(aₖ−y)²/(2σ²)] / ∑ⱼ₌₀⁸⁰ exp[−(aⱼ−y)²/(2σ²)]",
    2: "μ(p)=∑ₖ aₖpₖ,    v(p)=∑ₖ pₖ[aₖ−μ(p)]²\nμ̄=μ(p)/80,    H̄=−∑ₖ pₖln(pₖ+ε)/ln81,    V̄=v(p)/80²",
    3: "γ(p)=|∑ₖ pₖ{[aₖ−μ(p)]/√[v(p)+ε]}³|\nB(p)=∑ₖ₌₀² pₖ+∑ₖ₌₇₈⁸⁰ pₖ",
    4: "uᶜ = [GAP(F̃₃), s(sg(pᶜ)), eθ(sg(pᶜ))]",
    5: "α(g,i) = exp[z(g,i)] / ∑ⱼ₌₁³ exp[z(g,j)],    g = 1,…,G",
    6: "Fᶠ(g) = ∑ᵢ₌₁³ α(g,i)F̃(i,g)\nFᶠ = φ₁×₁(Concat_g Fᶠ(g))",
    7: "μ = ∑ₖ₌₀⁸⁰ aₖpᵐₖ",
    8: "dᵐ=[s(sg(pᵐ)),eψ(sg(pᵐ))]\ng=Sigmoid[gate(dᵐ)],    r=b tanh{res[GAP(Fᶠ),dᵐ]},    b=3",
    9: "ŷ = clip(μ + gr, 0, 80)",
    10: "ℒᶜ = D_KL(q ∥ pᶜ) + SL1(μᶜ,y)\nℒᵐ = D_KL(q ∥ pᵐ) + SL1(μ,y)",
    11: "g* = clip(|sg(μ)−y|/3,0,1)\nℒᵍ = SL1(g,g*),    ℒʳ = SL1(ŷ,y)",
    12: "ρ(t) = { 0,  t<16;    min[1,(t−16)/10],  t≥16 }",
    13: "ℒᴷᴰ = D_KL(pᵀ ∥ pᵐ) = ∑ₖpᵀₖ ln(pᵀₖ/pᵐₖ)",
    14: "ℒ = ℒᵐ + 0.3ℒᶜ + ρ(t)(0.5ℒʳ + 0.1ℒᵍ) + 1.0ℒᴷᴰ",
    15: "MAE = (1/N) ∑ₙ₌₁ᴺ |ŷₙ−yₙ|",
}


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_fonts(
    run,
    *,
    east_asia: str | None = None,
    latin: str | None = None,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    east_asia = east_asia or TOKENS["font"]["zh_body"]
    latin = latin or TOKENS["font"]["latin"]
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_fonts(style, *, east_asia: str, latin: str, size: float, bold: bool = False, color: str | None = None) -> None:
    style.font.name = latin
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def get_or_add_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_fonts(normal, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["body"], color=TOKENS["color"]["ink"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(TOKENS["spacing_pt"]["body_after"])
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_style_fonts(title, east_asia=TOKENS["font"]["zh_heading"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["title"], bold=True, color=TOKENS["color"]["ink"])
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.keep_with_next = True

    subtitle = document.styles["Subtitle"]
    set_style_fonts(subtitle, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["subtitle"], color=TOKENS["color"]["muted"])
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True

    h1 = document.styles["Heading 1"]
    set_style_fonts(h1, east_asia=TOKENS["font"]["zh_heading"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["h1"], bold=True, color=TOKENS["color"]["ink"])
    h1.paragraph_format.space_before = Pt(TOKENS["spacing_pt"]["h1_before"])
    h1.paragraph_format.space_after = Pt(TOKENS["spacing_pt"]["h1_after"])
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True
    h1.paragraph_format.page_break_before = False

    h2 = document.styles["Heading 2"]
    set_style_fonts(h2, east_asia=TOKENS["font"]["zh_heading"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["h2"], bold=True, color=TOKENS["color"]["accent"])
    h2.paragraph_format.space_before = Pt(TOKENS["spacing_pt"]["h2_before"])
    h2.paragraph_format.space_after = Pt(TOKENS["spacing_pt"]["h2_after"])
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    abstract_heading = get_or_add_style(document, "Abstract Heading")
    set_style_fonts(abstract_heading, east_asia=TOKENS["font"]["zh_heading"], latin=TOKENS["font"]["latin"], size=10.5, bold=True, color=TOKENS["color"]["accent"])
    abstract_heading.paragraph_format.space_before = Pt(6)
    abstract_heading.paragraph_format.space_after = Pt(3)
    abstract_heading.paragraph_format.keep_with_next = True

    abstract = get_or_add_style(document, "Abstract")
    set_style_fonts(abstract, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=9.5, color=TOKENS["color"]["ink"])
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.line_spacing = 1.2
    abstract.paragraph_format.first_line_indent = Pt(19)
    abstract.paragraph_format.space_after = Pt(3)

    keywords = get_or_add_style(document, "Keywords")
    set_style_fonts(keywords, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=9.5, color=TOKENS["color"]["ink"])
    keywords.paragraph_format.first_line_indent = Pt(0)
    keywords.paragraph_format.space_after = Pt(6)

    caption = document.styles["Caption"]
    set_style_fonts(caption, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["caption"], color=TOKENS["color"]["ink"])
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    table_caption = get_or_add_style(document, "Table Caption")
    set_style_fonts(table_caption, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["caption"], color=TOKENS["color"]["ink"])
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.first_line_indent = Pt(0)
    table_caption.paragraph_format.space_before = Pt(5)
    table_caption.paragraph_format.space_after = Pt(3)
    table_caption.paragraph_format.keep_with_next = True

    reference = get_or_add_style(document, "Reference")
    set_style_fonts(reference, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["reference"], color=TOKENS["color"]["ink"])
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.left_indent = Cm(0.75)
    reference.paragraph_format.first_line_indent = Cm(-0.75)
    reference.paragraph_format.line_spacing = 1.05
    reference.paragraph_format.space_after = Pt(1)

    list_number = document.styles["List Number"]
    set_style_fonts(list_number, east_asia=TOKENS["font"]["zh_body"], latin=TOKENS["font"]["latin"], size=TOKENS["size_pt"]["body"], color=TOKENS["color"]["ink"])
    list_number.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    list_number.paragraph_format.left_indent = Cm(0.8)
    list_number.paragraph_format.first_line_indent = Cm(-0.55)
    list_number.paragraph_format.space_after = Pt(3)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(TOKENS["page"]["margin_top_cm"])
    section.bottom_margin = Cm(TOKENS["page"]["margin_bottom_cm"])
    section.left_margin = Cm(TOKENS["page"]["margin_side_cm"])
    section.right_margin = Cm(TOKENS["page"]["margin_side_cm"])
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True


def add_table_bottom_rule(table, color: str, size: int) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_fonts(run, size=8.5, color=TOKENS["color"]["muted"])


def configure_headers_and_footers(document: Document) -> None:
    section = document.sections[0]
    width = section_content_width_dxa(section)

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    p = first_header.paragraphs[0]
    p._element.getparent().remove(p._element)
    table = first_header.add_table(rows=1, cols=2, width=Cm(16.2))
    apply_table_geometry(table, column_widths_from_weights([1.0, 1.0], width), table_width_dxa=width, indent_dxa=0, cell_margins_dxa={"top": 30, "bottom": 55, "start": 0, "end": 0})
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_fonts(left.add_run("RESEARCH ARTICLE"), size=8.2, bold=True, color=TOKENS["color"]["accent"])
    set_run_fonts(right.add_run("计算机视觉 · 人脸分析"), east_asia=TOKENS["font"]["zh_heading"], size=8.2, color=TOKENS["color"]["muted"])
    add_table_bottom_rule(table, TOKENS["color"]["accent"], 10)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(2)
    set_run_fonts(hp.add_run("FADE-Net：分布条件路由与有界残差细化"), east_asia=TOKENS["font"]["zh_body"], size=8.2, color=TOKENS["color"]["muted"])
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), TOKENS["color"]["rule"])
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(fp)


def add_horizontal_rule(paragraph, color: str = "4E708C", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def linearize_inline_math(text: str) -> str:
    replacements = (
        (r"\mathcal A", "𝒜"),
        (r"\mathcal L", "ℒ"),
        (r"\mathbb R", "ℝ"),
        (r"\operatorname{GAP}", "GAP"),
        (r"\operatorname{SL1}", "SL1"),
        (r"\operatorname{clip}", "clip"),
        (r"\operatorname{sg}", "sg"),
        (r"\mathrm{KL}", "KL"),
        (r"\hat y", "ŷ"),
        (r"\tilde F", "F̃"),
        (r"\mu", "μ"),
        (r"\sigma", "σ"),
        (r"\alpha", "α"),
        (r"\rho", "ρ"),
        (r"\theta", "θ"),
        (r"\psi", "ψ"),
        (r"\gamma", "γ"),
        (r"\varepsilon", "ε"),
        (r"\sqrt", "√"),
        (r"\rightarrow", "→"),
        (r"^{C\times H\times W}", "⁽ᶜ×ᴴ×ᵂ⁾"),
        (r"\times", "×"),
        (r"\in", "∈"),
        (r"\ldots", "…"),
        (r"\|", "∥"),
    )
    for source, target in replacements:
        text = text.replace(source, target)
    superscript_map = str.maketrans("-0123456789", "⁻⁰¹²³⁴⁵⁶⁷⁸⁹")
    text = re.sub(r"\^\{(-?\d+)\}", lambda match: match.group(1).translate(superscript_map), text)
    for source, target in (
        ("^T", "ᵀ"),
        ("^c", "ᶜ"),
        ("^m", "ᵐ"),
        ("^f", "ᶠ"),
        ("^5", "⁵"),
        ("_1", "₁"),
        ("_2", "₂"),
        ("_3", "₃"),
        ("_c", "ᶜ"),
        ("_k", "ₖ"),
        ("_i", "ᵢ"),
        ("_j", "ⱼ"),
        ("_n", "ₙ"),
    ):
        text = text.replace(source, target)
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    text = text.replace("_θ", "θ").replace("_ψ", "ψ")
    return text


def add_rich_text(
    paragraph,
    text: str,
    *,
    style_size: float | None = None,
    format_citations: bool = True,
) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            set_run_fonts(paragraph.add_run(text[position : match.start()]), size=style_size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_fonts(run, east_asia="等线", latin="Consolas", size=(style_size or TOKENS["size_pt"]["body"]) - 0.5, color=TOKENS["color"]["muted"])
        elif token.startswith("$"):
            run = paragraph.add_run(linearize_inline_math(token[1:-1]))
            set_run_fonts(run, east_asia=TOKENS["font"]["math"], latin=TOKENS["font"]["math"], size=style_size, italic=True)
        else:
            run = paragraph.add_run(token)
            set_run_fonts(run, size=style_size)
            run.font.superscript = CITATIONS_AS_SUPERSCRIPT if format_citations else False
        position = match.end()
    if position < len(text):
        set_run_fonts(paragraph.add_run(text[position:]), size=style_size)


def add_omath(paragraph, text: str) -> None:
    lines = text.split("\n")
    size = 9.0 if max(map(len, lines)) > 85 else 9.5
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_fonts(
            run,
            east_asia=TOKENS["font"]["math"],
            latin=TOKENS["font"]["math"],
            size=size,
        )
        if index + 1 < len(lines):
            run.add_break(WD_BREAK.LINE)


def add_equation(document: Document, latex_lines: list[str], tag: int) -> None:
    section = document.sections[0]
    width = section_content_width_dxa(section)
    table = document.add_table(rows=1, cols=3)
    apply_table_geometry(table, column_widths_from_weights([1.2, 7.6, 1.2], width), table_width_dxa=width, indent_dxa=0, cell_margins_dxa={"top": 35, "bottom": 35, "start": 0, "end": 0})
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"}, insideH={"val": "nil"}, insideV={"val": "nil"})
    equation_p = table.cell(0, 1).paragraphs[0]
    equation_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_p.paragraph_format.space_before = Pt(2)
    equation_p.paragraph_format.space_after = Pt(2)
    equation_p.paragraph_format.keep_together = True
    add_omath(equation_p, EQUATION_LINEAR.get(tag, " ".join(latex_lines)))
    number_p = table.cell(0, 2).paragraphs[0]
    number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_p.paragraph_format.keep_together = True
    set_run_fonts(number_p.add_run(f"（{tag}）"), size=9)
    prevent_row_split(table.rows[0])


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_weights(caption: str, columns: int) -> list[float]:
    if caption.startswith("表1") and columns == 5:
        return [1.0, 2.6, 1.7, 1.7, 1.2]
    if caption.startswith("表2") and columns == 7:
        return [2.5, 1.0, 1.0, 1.0, 1.0, 1.0, 1.8]
    if caption.startswith("表3") and columns == 4:
        return [1.8, 2.8, 2.5, 1.3]
    if caption.startswith("表4") and columns == 4:
        return [1.7, 2.7, 1.4, 1.4]
    if caption.startswith("表5") and columns == 6:
        return [1.7, 1.1, 1.4, 1.4, 1.6, 1.8]
    return [1.0] * columns


def clean_table_text(text: str) -> str:
    text = re.sub(r"\$([^$]+)\$", lambda m: linearize_inline_math(m.group(1)), text)
    return text.replace("\\", "")


def add_markdown_table(document: Document, caption: str, table_lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in table_lines if not TABLE_SEPARATOR_RE.match(line)]
    if not rows:
        return
    columns = len(rows[0])
    if any(len(row) != columns for row in rows):
        raise ValueError(f"inconsistent Markdown table after {caption}")
    section = document.sections[0]
    width = section_content_width_dxa(section)
    table = document.add_table(rows=len(rows), cols=columns)
    apply_table_geometry(
        table,
        column_widths_from_weights(table_weights(caption, columns), width),
        table_width_dxa=width,
        indent_dxa=70,
        cell_margins_dxa={"top": 55, "bottom": 55, "start": 70, "end": 70},
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, (word_row, values) in enumerate(zip(table.rows, rows)):
        prevent_row_split(word_row)
        for col_idx, (cell, value) in enumerate(zip(word_row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(clean_table_text(value))
            set_run_fonts(run, size=TOKENS["size_pt"]["table"], bold=row_idx == 0)
            borders: dict[str, dict[str, str]] = {
                "left": {"val": "nil"},
                "right": {"val": "nil"},
                "insideH": {"val": "nil"},
                "insideV": {"val": "nil"},
            }
            if row_idx == 0:
                borders["top"] = {"val": "single", "sz": "12", "color": TOKENS["color"]["ink"]}
                borders["bottom"] = {"val": "single", "sz": "6", "color": TOKENS["color"]["ink"]}
                shade_cell(cell, TOKENS["color"]["header_fill"])
            elif row_idx == len(rows) - 1:
                borders["bottom"] = {"val": "single", "sz": "12", "color": TOKENS["color"]["ink"]}
            else:
                borders["top"] = {"val": "nil"}
                borders["bottom"] = {"val": "nil"}
            set_cell_border(cell, **borders)
    set_repeat_table_header(table.rows[0])
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.first_line_indent = Pt(0)


def add_picture(document: Document, alt: str, relative_path: str) -> None:
    image_path = (SOURCE.parent / relative_path).resolve()
    if not image_path.is_file():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    width_cm = 13.5 if image_path.name.startswith("fig3_") else 16.0
    shape = run.add_picture(str(image_path), width=Cm(width_cm))
    shape._inline.docPr.set("descr", alt)


def add_body_paragraph(document: Document, text: str, *, abstract_mode: bool = False) -> None:
    if text.startswith("关键词：") or text.startswith("Key words:"):
        paragraph = document.add_paragraph(style="Keywords")
        label, value = text.split("：", 1) if "：" in text else text.split(":", 1)
        run = paragraph.add_run(label + ("：" if "：" in text else ":"))
        set_run_fonts(run, east_asia=TOKENS["font"]["zh_heading"], size=9.5, bold=True)
        add_rich_text(paragraph, " " + value.strip(), style_size=9.5)
        return
    style = "Abstract" if abstract_mode else "Normal"
    paragraph = document.add_paragraph(style=style)
    add_rich_text(paragraph, text, style_size=9.5 if abstract_mode else TOKENS["size_pt"]["body"])
    if text.startswith("五折所有结果文件的数据划分指纹"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)


def add_title_block(document: Document, title: str, subtitle: str) -> None:
    title_p = document.add_paragraph(style="Title")
    add_rich_text(title_p, title, style_size=TOKENS["size_pt"]["title"])
    subtitle_p = document.add_paragraph(style="Subtitle")
    add_rich_text(subtitle_p, subtitle, style_size=TOKENS["size_pt"]["subtitle"])
    rule = document.add_paragraph()
    rule.paragraph_format.first_line_indent = Pt(0)
    rule.paragraph_format.space_after = Pt(5)
    add_horizontal_rule(rule, TOKENS["color"]["accent"], 8)


def build_document() -> Document:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    if len(lines) < 3 or not lines[0].startswith("# ") or not lines[2].startswith("## "):
        raise ValueError("unexpected manuscript title block")
    document = Document()
    configure_page(document)
    configure_styles(document)
    configure_headers_and_footers(document)
    add_title_block(document, lines[0][2:].strip(), lines[2][3:].strip())

    document.core_properties.title = lines[0][2:].strip()
    document.core_properties.subject = "轻量级人脸年龄估计；AFAD五个官方主体互斥划分"
    document.core_properties.keywords = "FADE-Net; DCSR; CGBR; AFAD; facial age estimation"
    document.core_properties.comments = f"Preset={PRESET}; header={HEADER_PATTERN}; generated from {SOURCE.name}"

    index = 3
    abstract_mode = False
    pending_table_caption: str | None = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line == "$$":
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation_lines.append(lines[index].strip())
                index += 1
            equation_text = " ".join(equation_lines)
            match = re.search(r"\\tag\{(\d+)\}", equation_text)
            if not match:
                raise ValueError(f"equation without tag: {equation_text}")
            add_equation(document, equation_lines, int(match.group(1)))
            index += 1
            continue
        figure = FIGURE_RE.match(line)
        if figure:
            add_picture(document, figure.group("alt"), figure.group("path"))
            index += 1
            continue
        if line.startswith("图") and re.match(r"^图\d+\s", line):
            caption = document.add_paragraph(style="Caption")
            add_rich_text(caption, line, style_size=TOKENS["size_pt"]["caption"])
            index += 1
            continue
        if line.startswith("表") and re.match(r"^表\d+\s", line):
            pending_table_caption = line
            caption = document.add_paragraph(style="Table Caption")
            add_rich_text(caption, line, style_size=TOKENS["size_pt"]["caption"])
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines) and lines[lookahead].strip().startswith("|"):
                table_lines: list[str] = []
                index = lookahead
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                add_markdown_table(document, pending_table_caption, table_lines)
                pending_table_caption = None
            else:
                index += 1
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            if text in {"摘要", "Abstract"}:
                abstract_mode = True
                paragraph = document.add_paragraph(style="Abstract Heading")
                add_rich_text(paragraph, text, style_size=10.5)
            else:
                abstract_mode = False
                paragraph = document.add_paragraph(style="Heading 2")
                add_rich_text(paragraph, text, style_size=TOKENS["size_pt"]["h2"])
            index += 1
            continue
        if line.startswith("## "):
            abstract_mode = False
            paragraph = document.add_paragraph(style="Heading 1")
            add_rich_text(paragraph, line[3:].strip(), style_size=TOKENS["size_pt"]["h1"])
            index += 1
            continue
        numbered = NUMBERED_RE.match(line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, numbered.group("text"), style_size=TOKENS["size_pt"]["body"])
            index += 1
            continue
        reference = REFERENCE_RE.match(line)
        if reference:
            paragraph = document.add_paragraph(style="Reference")
            add_rich_text(
                paragraph,
                f"[{reference.group('number')}] {reference.group('text')}",
                style_size=TOKENS["size_pt"]["reference"],
                format_citations=False,
            )
            index += 1
            continue
        add_body_paragraph(document, line, abstract_mode=abstract_mode)
        if line.startswith("Key words:"):
            abstract_mode = False
        index += 1

    return document


def main() -> None:
    document = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")
    print(f"preset={PRESET}; header_pattern={HEADER_PATTERN}")


if __name__ == "__main__":
    main()
