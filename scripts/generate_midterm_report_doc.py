from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT_DIR / "output" / "doc"
OUTPUT_DOCX = OUTPUT_DIR / "基于注意力机制的轻量化人脸年龄估计_中期汇报发言稿与QA.docx"


def set_run_font(run, font_name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14)]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("基于注意力机制的轻量化人脸年龄估计\n中期汇报发言稿与问答模板")
    set_run_font(run, font_name="黑体", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("以下内容可直接用于中期汇报口头陈述，并可根据现场时间适当删减。")
    set_run_font(run, size=11)


def add_quick_facts(doc: Document) -> None:
    doc.add_heading("一、汇报速记信息", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("研究题目", "基于注意力机制的轻量化人脸年龄估计"),
        ("数据集与划分", "AFAD，分层 72-8-20；训练/验证/测试分别为 117603 / 13039 / 32731"),
        ("核心模型", "MobileNetV3-Large + Coordinate Attention + 双流融合 + Bottleneck SPP"),
        ("最佳结果", "历史最佳单模型 MAE 为 3.0574（seed1337，需按当前协议复核）"),
        ("稳定性结果", "历史三次实验均值为 3.0857 ± 0.0204，正式论文需核验日志与元数据"),
        ("参数规模", "总参数量为 4.8415M"),
    ]
    for i, (left, right) in enumerate(rows):
        for j, text in enumerate((left, right)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=11, bold=(j == 0))


def add_speech(doc: Document) -> None:
    doc.add_heading("二、中期汇报简要发言稿", level=1)
    paragraphs = [
        "各位老师好，我汇报的题目是《基于注意力机制的轻量化人脸年龄估计》。本课题主要关注在保证模型轻量化和可部署性的前提下，提高人脸年龄估计的准确性与稳定性。之所以选择这一方向，是因为年龄估计在智慧零售、安防分析、人机交互等场景中具有较强的应用价值，但现有高精度方法往往依赖较重的网络结构，不利于资源受限环境部署。",
        "目前，我已经完成了相关文献梳理、整体技术路线设计、数据集整理与核心模型搭建。方法上以 MobileNetV3-Large 作为轻量骨干，在深层阶段将最后 4 个含有 SE 的模块替换为 Coordinate Attention，以增强对关键面部区域的空间敏感性；同时从中间层提取纹理分支与语义分支进行双流融合，并在深层加入 Bottleneck SPP，用较小的额外代价增强多尺度上下文建模能力。",
        "在监督目标和训练策略上，我采用了标签分布学习框架，并结合 L1 约束、CDF ranking 损失以及 Mean-Variance loss，以缓解年龄标签模糊、相邻年龄难分和序关系表达不足的问题。实验数据集使用 AFAD，当前统一采用分层 72-8-20 的划分协议。训练过程中前 10 轮冻结骨干网络，随后进行全参数微调，并配合 AdamW、EMA、MixUp 和测试时增强等策略提高收敛稳定性。",
        "从已有历史记录来看，模型在固定协议下具备较好的精度与参数量平衡。历史最佳单模型在 seed1337 下的测试集 MAE 为 3.0574，另外两个随机种子的记录分别为 3.0951 和 3.1047，三次实验均值为 3.0857，标准差为 0.0204。正式论文中，这些数字需要与当前 split、TTA、checkpoint 元数据和日志一起复核后再作为主表依据。",
        "下一阶段，我准备重点完成三项工作。第一，补充更系统的消融实验，进一步验证注意力注入、双流融合、SPP 和损失项组合的独立贡献；第二，增加注意力可视化与失败样本分析，提升方法解释性；第三，继续打磨论文文字表达与结果展示，使整体论证更加完整。我的汇报完毕，请各位老师批评指正。",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, size=12)


def add_qa(doc: Document) -> None:
    doc.add_heading("三、高频问答模板", level=1)
    qa_items = [
        (
            "Q1 这个课题的研究意义是什么？",
            "答题模板：这个课题的意义主要在于两点。第一，人脸年龄估计在智慧零售、安防分析和人机交互中都有明确应用场景；第二，现有高精度方法大多依赖较重模型，不利于移动端和边缘端部署。我的工作希望在精度和轻量化之间找到更合理的平衡。"
        ),
        (
            "Q2 为什么选择 MobileNetV3-Large 作为骨干网络？",
            "答题模板：选择 MobileNetV3-Large，主要是因为它在轻量化网络中具有较好的成熟度和工程稳定性，参数规模和推理成本都更适合部署场景。同时，它本身保留了一定的表征能力，便于在其基础上叠加任务定制模块，而不是从零设计一个新的骨干。"
        ),
        (
            "Q3 你的方法相对于基线模型，核心改进点是什么？",
            "答题模板：核心改进点有四个。第一，在深层按层注入 Coordinate Attention；第二，引入 Block 6 和 Block 12 的双流融合结构；第三，在深层特征后加入 Bottleneck SPP；第四，在训练目标上结合标签分布学习、L1、排序损失和均值方差约束。整体思路是围绕年龄估计任务特性，对轻量骨干进行有针对性的增强。"
        ),
        (
            "Q4 为什么只替换最后 4 个 SE 模块，而不是全部替换？",
            "答题模板：我的考虑是浅层特征更多对应边缘和基础纹理，统一替换的收益有限，反而可能带来额外开销。深层特征更接近面部语义表达，对关键区域的位置敏感性要求更高，因此只替换最后 4 个模块，能够在成本可控的前提下提升有效性。"
        ),
        (
            "Q5 为什么不用纯回归，而要采用标签分布学习和排序损失？",
            "答题模板：年龄标签具有连续性和模糊性，相邻年龄之间往往没有非常清晰的边界。纯回归虽然简单，但对这种模糊性表达不足。标签分布学习能够描述邻近年龄的概率关系，排序损失则可以进一步强化年龄的有序性，因此更符合任务本身特点。"
        ),
        (
            "Q6 为什么选择 AFAD 数据集，并采用 72-8-20 的划分方式？",
            "答题模板：AFAD 的样本规模较大，年龄标签相对明确，适合开展年龄估计实验。至于 72-8-20 的划分，是因为当前代码配置、结果文件和实验流程都统一基于这一协议，并且采用分层划分后能较好保持各年龄段样本比例一致，便于结果比较与复现。"
        ),
        (
            "Q7 目前阶段最能说明方法有效性的结果是什么？",
            "答题模板：目前最直接的历史记录有三个。第一，最佳单模型 MAE 记录为 3.0574；第二，三次不同随机种子的历史均值为 3.0857，标准差为 0.0204；第三，模型参数量约为 4.8415M。正式论文中，这些结果需要和当前 split、TTA、checkpoint 元数据及日志一起复核后再作为主表依据。"
        ),
        (
            "Q8 你现在这个阶段还存在哪些不足？",
            "答题模板：目前主要有三点不足。第一，高龄样本分布偏少，尾部年龄段的预测误差可能更大；第二，完整消融实验还需要继续补充；第三，注意力可视化和跨数据集验证还不够充分。下一阶段我会优先补齐这些部分，使论证更加完整。"
        ),
        (
            "Q9 如果老师认为创新点不够强，你怎么回应？",
            "答题模板：我会把创新点界定为面向具体任务的结构优化，而不是完全提出全新骨干。这个工作的价值在于结合轻量化部署需求，把注意力注入、多尺度融合和分布学习有机整合到同一框架中，并通过真实实验结果证明这种组合在参数受限条件下仍然有效。"
        ),
        (
            "Q10 后续准备怎么推进，确保按时完成？",
            "答题模板：后续我会按三个方向推进。第一，补充消融实验与可视化分析；第二，进一步整理结果表格、训练曲线和关键图示；第三，继续完善论文结构和文字表达。我的目标是先把实验论证补完整，再集中打磨文字与版式，这样可以保证内容和形式同步收口。"
        ),
    ]
    for question, answer in qa_items:
        doc.add_heading(question, level=2)
        p = doc.add_paragraph()
        run = p.add_run(answer)
        set_run_font(run, size=12)


def build_doc() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_title(doc)
    add_quick_facts(doc)
    add_speech(doc)
    add_qa(doc)
    doc.core_properties.title = "基于注意力机制的轻量化人脸年龄估计中期汇报发言稿与问答模板"
    doc.core_properties.subject = "中期汇报材料"
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_doc()
    print(output)
