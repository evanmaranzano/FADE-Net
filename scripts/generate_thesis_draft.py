from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT_DIR = Path(__file__).resolve().parent.parent
SOURCE_MD = ROOT_DIR / "docs" / "thesis_initial_draft.md"
OUTPUT_DIR = ROOT_DIR / "output" / "doc"
OUTPUT_DOCX = OUTPUT_DIR / "基于注意力机制的轻量化人脸年龄估计.docx"
ASSET_DIR = OUTPUT_DIR / "assets"


def set_run_font(run, font_name="宋体", size=12, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "黑体"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.first_line_indent = Cm(0)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "黑体"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.first_line_indent = Cm(0)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "黑体"
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(2)
    heading3.paragraph_format.first_line_indent = Cm(0)


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("学术论文")
    set_run_font(run, font_name="黑体", size=20, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于注意力机制的轻量化人脸年龄估计")
    set_run_font(run, font_name="黑体", size=22, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    fields = [
        "学校：____________________________",
        "学院：____________________________",
        "专业：____________________________",
        "学生姓名：________________________",
        "学号：____________________________",
        "指导教师：________________________",
        "完成日期：________________________",
    ]
    for item in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(item)
        set_run_font(run, size=14)

    add_page_break(doc)


def add_statement(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("诚信声明")
    set_run_font(run, font_name="黑体", size=18, bold=True)

    paragraphs = [
        "本人郑重声明：本文是在相关研究、实验分析与文献梳理基础上独立完成的学术论文。除文中已经明确注明引用的内容外，不包含他人已经发表或撰写的研究成果，也未以任何形式重复提交。",
        "本人承诺严格遵守学术规范，对文中所使用的资料、数据、图表与研究结论负责，并按照相关要求进一步完善文稿格式与引文标注。",
        "声明人（签名）：________________________",
        "日期：_______________________________",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if text.startswith("声明人") or text.startswith("日期"):
            p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_run_font(run, size=12)

    add_page_break(doc)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("目录")
    set_run_font(run, font_name="黑体", size=18, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "右键更新目录或按 F9 更新"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_separate)
    r._r.append(fld_text)
    r._r.append(fld_end)
    set_run_font(r, size=12)
    add_page_break(doc)


def strip_inline_markup(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue
        cells = [strip_inline_markup(cell) for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(cell_text)
            set_run_font(run, size=11, bold=(i == 0))


def add_image(doc: Document, alt_text: str, image_path: Path) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(15.5))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    run = caption.add_run(alt_text)
    set_run_font(run, size=11)


def load_font(size: int, bold: bool = False):
    candidates = []
    if bold:
        candidates.extend(
            [
                "C:/Windows/Fonts/msyhbd.ttc",
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/arialbd.ttf",
            ]
        )
    candidates.extend(
        [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/arial.ttf",
        ]
    )
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text, font, fill="#1F1F1F", spacing=8):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = left + (right - left - text_w) / 2
    y = top + (bottom - top - text_h) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=spacing, align="center")


def generate_architecture_diagram(output_path: Path) -> None:
    width, height = 1800, 1400
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    text_font = load_font(28)

    def box(x1, y1, x2, y2, text, color):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=color, outline="#2F4F4F", width=4)
        draw_centered_text(draw, (x1 + 18, y1 + 18, x2 - 18, y2 - 18), text, text_font)

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#555555", width=6)
        dx, dy = x2 - x1, y2 - y1
        if abs(dx) > abs(dy):
            sign = 1 if dx >= 0 else -1
            head = [(x2, y2), (x2 - 24 * sign, y2 - 12), (x2 - 24 * sign, y2 + 12)]
        else:
            sign = 1 if dy >= 0 else -1
            head = [(x2, y2), (x2 - 12, y2 - 24 * sign), (x2 + 12, y2 - 24 * sign)]
        draw.polygon(head, fill="#555555")

    draw_centered_text(draw, (120, 40, width - 120, 120), "FADE-Net Architecture Derived from Local Code", title_font)

    box(560, 150, 1240, 270, "Input Face Image\n224 x 224 x 3", "#F9E79F")
    box(420, 340, 1380, 470, "MobileNetV3-Large Backbone", "#D6EAF8")
    box(80, 560, 650, 700, "Block 6\nTexture Branch (40ch)", "#E8DAEF")
    box(1150, 560, 1720, 700, "Block 12\nSemantic Branch (112ch)", "#E8DAEF")
    box(420, 760, 1380, 880, "Last 4 SE Blocks -> CoordAtt", "#D5F5E3")
    box(420, 940, 1380, 1060, "Bottleneck SPP\n5 / 9 / 13 MaxPool -> 512ch", "#FADBD8")
    box(420, 1120, 1380, 1240, "Feature Fusion\nTexture + Semantic + SPP", "#FCF3CF")
    box(520, 1285, 1280, 1380, "Prediction Head  ->  1024  ->  81  ->  Expected Age", "#D6DBDF")

    arrow(900, 270, 900, 340)
    arrow(900, 470, 360, 560)
    arrow(900, 470, 1440, 560)
    arrow(360, 700, 760, 760)
    arrow(1440, 700, 1040, 760)
    arrow(900, 880, 900, 940)
    arrow(900, 1060, 900, 1120)
    arrow(900, 1240, 900, 1285)

    image.save(output_path)


def generate_age_distribution_chart(output_path: Path) -> None:
    dataset_dir = ROOT_DIR / "datasets" / "AFAD"
    ages = []
    counts = []
    for age_dir in sorted(dataset_dir.iterdir(), key=lambda p: int(p.name) if p.name.isdigit() else 10**9):
        if age_dir.is_dir() and age_dir.name.isdigit():
            count = sum(1 for f in age_dir.rglob("*") if f.suffix.lower() in {".jpg", ".png"})
            if count > 0:
                ages.append(int(age_dir.name))
                counts.append(count)

    width, height = 1800, 900
    margin_left, margin_right = 120, 60
    margin_top, margin_bottom = 120, 110
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(38, bold=True)
    text_font = load_font(24)
    small_font = load_font(20)

    draw_centered_text(draw, (100, 25, width - 100, 95), "AFAD Age Distribution (Local Processed Dataset)", title_font)
    draw.line((margin_left, margin_top, margin_left, margin_top + chart_h), fill="#444444", width=3)
    draw.line(
        (margin_left, margin_top + chart_h, margin_left + chart_w, margin_top + chart_h),
        fill="#444444",
        width=3,
    )

    max_count = max(counts)
    tick_count = 5
    for i in range(tick_count + 1):
        value = max_count * i / tick_count
        y = margin_top + chart_h - chart_h * i / tick_count
        draw.line((margin_left, y, margin_left + chart_w, y), fill="#D5D8DC", width=1)
        draw.text((30, y - 12), f"{int(value):,}", font=small_font, fill="#4D4D4D")

    bar_w = chart_w / max(len(ages), 1)
    for idx, (age, count) in enumerate(zip(ages, counts)):
        x1 = margin_left + idx * bar_w + 2
        x2 = margin_left + (idx + 1) * bar_w - 2
        y2 = margin_top + chart_h
        y1 = y2 - chart_h * count / max_count
        draw.rectangle((x1, y1, x2, y2), fill="#5DADE2", outline="#2E86C1")
        if idx % 5 == 0 or idx == len(ages) - 1:
            label = str(age)
            bbox = draw.textbbox((0, 0), label, font=small_font)
            text_x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
            draw.text((text_x, y2 + 12), label, font=small_font, fill="#4D4D4D")

    draw.text((width // 2 - 30, height - 52), "Age", font=text_font, fill="#2F2F2F")
    draw.text((margin_left - 85, margin_top - 55), "Samples", font=text_font, fill="#2F2F2F")
    image.save(output_path)


def prepare_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    generate_architecture_diagram(ASSET_DIR / "fade_architecture.png")
    generate_age_distribution_chart(ASSET_DIR / "afad_age_distribution.png")


def add_body_paragraph(doc: Document, text: str) -> None:
    text = strip_inline_markup(text)
    if not text:
        return

    p = doc.add_paragraph()
    if text.startswith("图 ") or text.startswith("表 ") or text.startswith("关键词") or text.startswith("Keywords"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.startswith(("图 ", "表 ")) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0 if text.startswith(("图 ", "表 ")) else 0.74)
    elif re.match(r"^\d+\.", text):
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_formula(doc: Document, formula_lines: list[str]) -> None:
    formula = "\n".join(line.strip() for line in formula_lines if line.strip())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(formula)
    set_run_font(run, font_name="Times New Roman", size=11)


def render_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    idx = 0
    paragraph_buffer: list[str] = []
    first_heading = True

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if not stripped:
            flush_paragraph_buffer()
            idx += 1
            continue

        if stripped == "---":
            flush_paragraph_buffer()
            idx += 1
            continue

        if stripped == "$$":
            flush_paragraph_buffer()
            idx += 1
            formula_lines = []
            while idx < len(lines) and lines[idx].strip() != "$$":
                formula_lines.append(lines[idx])
                idx += 1
            add_formula(doc, formula_lines)
            idx += 1
            continue

        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph_buffer()
            alt_text, raw_path = image_match.groups()
            add_image(doc, alt_text, (ROOT_DIR / raw_path).resolve())
            idx += 1
            continue

        if is_table_line(stripped):
            flush_paragraph_buffer()
            table_lines = []
            while idx < len(lines) and is_table_line(lines[idx]):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(doc, parse_table(table_lines))
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph_buffer()
            level = len(heading_match.group(1))
            title = strip_inline_markup(heading_match.group(2))
            if title == "目录":
                add_toc(doc)
                first_heading = False
                idx += 1
                continue
            if level == 1 and not first_heading:
                add_page_break(doc)
            style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
            p = doc.add_paragraph(style=style_name)
            if title in {"摘要", "Abstract", "参考文献", "致谢"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(title)
            set_run_font(
                run,
                font_name="黑体" if level <= 2 else "黑体",
                size=16 if level == 1 else 14 if level == 2 else 12,
                bold=True,
            )
            first_heading = False
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph_buffer()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "基于注意力机制的轻量化人脸年龄估计"
    doc.core_properties.subject = "学术论文"
    doc.core_properties.author = ""

    add_cover(doc)
    add_statement(doc)
    render_markdown(doc, SOURCE_MD.read_text(encoding="utf-8"))
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    prepare_assets()
    doc = build_document()
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
